import os
import io
import csv
import docx # type: ignore
import logging
import google.generativeai as genai # type: ignore

from datetime import datetime
from google.oauth2 import service_account # type: ignore
from googleapiclient.discovery import build # type: ignore
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload # type: ignore

SCOPES = ['https://www.googleapis.com/auth/drive']
TO_BE_PROCESSED_FOLDER_ID = 'REMOVED'
PROCESSED_FOLDER_ID = 'REMOVED'
OUTPUT_FOLDER_ID = 'REMOVED'
LOGGING_FOLDER_ID = 'REMOVED'
TIMESTAMP = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
LOG_FILENAME = f"{TIMESTAMP}_parser.log"
OUTPUT_CSV_FILENAME = f"{TIMESTAMP}_output.csv"

def config_logger():
  logger = logging.getLogger()
  logger.setLevel(logging.DEBUG)
  formatter = logging.Formatter('%(levelname)s:%(asctime)s: %(message)s')
  file_handler = logging.FileHandler(LOG_FILENAME)
  file_handler.setLevel(logging.INFO)
  file_handler.setFormatter(formatter)
  logger.addHandler(file_handler)
  logging.getLogger('googleapiclient.discovery_cache').setLevel(logging.ERROR)
  return logger

logger = config_logger()

def get_google_drive_service():
  creds: None
  if os.path.exists('local_service_account.json'):
    logger.info('service_account.json file found! Creating credentials...')
    creds = service_account.Credentials.from_service_account_file('local_service_account.json', scopes=SCOPES)
  return build('drive', 'v3', credentials=creds)

def get_file_list(service):
  full_results = []
  results = service.files().list(q=f"'{TO_BE_PROCESSED_FOLDER_ID}' in parents", pageSize=1000, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
  next_page_token = results.get('nextPageToken')
  full_results = (results.get('files', []))
  
  while (next_page_token != None):
      results = service.files().list(q=f"'{TO_BE_PROCESSED_FOLDER_ID}' in parents", pageSize=1000, pageToken=next_page_token, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
      next_page_token = results.get('nextPageToken')
      full_results.extend(results.get('files', []))
  
  return full_results

def read_word_file_from_drive(service, file_id):
  request = service.files().get_media(fileId=file_id)
  file_stream = io.BytesIO()
  downloader = MediaIoBaseDownload(file_stream, request)
  done = False
  while not done:
    done = downloader.next_chunk()
  file_stream.seek(0)
  return file_stream


def extract_text(document):
  text = []
  for section in document.sections:
    for paragraph in section.header.paragraphs:
      text.append(paragraph.text)
    for paragraph in section.footer.paragraphs:
      text.append(paragraph.text)
  for paragraph in document.paragraphs:
    text.append(paragraph.text)
  for table in document.tables:
    for row in table.rows:
      for cell in row.cells:
        text.append(cell.text)
  return text
  
    
def call_gemini(prompt_text):
  genai.configure(api_key='REMOVED')
  model = genai.GenerativeModel('gemini-2.5-flash')
  response = model.generate_content(
    "Task: You will be given the full text of an agreement. Extract every milestone obligation and output them only as CSV rows with the following six columns, in this order:\n"
    "AgreementNumber — the value that follows the literal text “OTC Agreement Number:” (e.g., A-2025-2025).\n"
    "MilestoneName — one of exactly two strings:\n"
    "- PERFORMANCE MILESTONES for obligations listed in Section 9 (Performance Milestones) or described with phrasing like “Licensee will complete milestone ,,,”.\n"
    "- MILESTONE PAYMENTS for obligations listed in Section 10.5 (Milestone Payments) or described with phrasing like 'Upon ...' or 'Licensee shall pay ...'.\n"
    "MilestoneTargetDate — the due date for the obligation in YYYY-MM-DD format.\n"
    "- If the milestone has no explicit date, leave this field blank.\n"
    "MilestoneDescription — For each milestone (e.g., a, b, c), include ALL sentences in that item, from the first sentence starting with 'Upon ...' or 'Licensee shall pay ...' to the last sentence **before the next item or section**. Wrap the entire block in double quotes. Do not shorten or stop early.\n"
    "- These will only ever be found inside sections titled 'Performance Milestones' and 'Milestone Payments'.\n"
    "- If the final character in the milestone description is a semicolon, remove it. Do not include trailing semicolons.\n"
    "- Only extract milestones from the text inside Sections 9 and 10.5. Ignore all other sections and text.\n"
    "- There may be instances where after the last performance milestone or milestone payment, there is an extra description. Ignore this."
    "MilestoneSetDeadline — TRUE if a due date exists for the milestone, otherwise FALSE.\n"
    "MilestonePayment — the numeric dollar amount associated with the milestone, with no currency symbol or commas (e.g., 3610).\n"
    "- If no payment is specified, leave this field blank.\n"
    "Formatting rules:\n"
    "- Output only CSV text.\n"
    "- Use a single comma (\",\") as the delimiter between each field.\n"
    "- Preserve the order in which milestones appear in the agreement.\n"
    "- Do NOT include a header row. Begin immediately with the first milestone row.\n"
    "Input: " + prompt_text
  , generation_config={"temperature": 0})
  return response


def upload_file_to_drive(service, folder_id, file_name, mime_type):
  file_metadata = {
    'name': file_name,
    'parents': [folder_id]
  }
  media = MediaFileUpload(file_name, mimetype=mime_type)
  file = service.files().create(body=file_metadata, media_body=media, fields='id', supportsAllDrives=True).execute()
  print(f"Uploaded File ID: {file.get('id')}; File Name: {file_name}")
  os.remove(file_name)
    
    
def move_file_to_processed(service, file_id):
  try:
    file = service.files().get(fileId=file_id, fields='parents', supportsAllDrives=True).execute()
    previous_parents = ",".join(file.get('parents'))
    file = service.files().update(
      fileId=file_id,
      addParents=PROCESSED_FOLDER_ID,
      removeParents=previous_parents,
      fields='id, parents',
      supportsAllDrives=True
    ).execute()
  except Exception as e:
    logger.error(f'Error moving file to \'Processed\' folder: {str(e)}')
    raise


def parse_documents(file_list, service):
  logger.info('Creating output.csv file...')
  with open(OUTPUT_CSV_FILENAME, 'w', newline='') as csvfile:
    writer = csv.writer(csvfile)
    logger.info('Writing header row in output.csv')
    writer.writerow([
      "OTC Agreement Number",
      "Milestone Name",
      "Milestone Target Completion Date",
      "Milestone Description",
      "Milestone Set Deadline",
      "Milestone Payment"
    ])
    logger.info('Beginning file parsing...')
    for file in file_list:
      file_id = file['id']
      file_name = file['name']
      logger.info(f'Reading {file_name} from Google Drive...')
      print(f'Processing {file_name}')
      word_stream = read_word_file_from_drive(service, file_id)
      document = docx.Document(word_stream)
      logger.info(f'Extracting text from {file_name}...')
      document_text = extract_text(document)
      logger.info(f'Creating prompt text for {file_name}...')
      prompt_text = "\n".join(document_text)
      logger.info('Calling Gemini API...')
      output_raw_text = call_gemini(prompt_text)
      logger.info('Formatting output text...')
      output_text = ''.join(part.text for part in output_raw_text.candidates[0].content.parts)
      file_object = io.StringIO(output_text)
      logger.info(f'Writing milestones from {file_name} into output.csv')
      reader = csv.reader(file_object)
      for row in reader:
        writer.writerow(row)
      logger.info(f'Successfully parsed milestones from {file_name}')
      print(f'Successfully parsed milestones from {file_name}')
      
    
def main():
  try:
    logger.info('Starting parser...')
    service = get_google_drive_service()
    logger.info('Credentials obtained.')
    file_list = get_file_list(service)
    logger.info('"To Be Processed" file list obtained.')
        
    if not file_list:
      logger.warning('No files found in \'To Be Processed\' folder.')
      return
    
    parse_documents(file_list, service)
    logger.info('Shutting down parser...')
  except Exception as e:
    logger.error(f"An unexpected error occurred: {e}")
    
if __name__ == '__main__':
    main()