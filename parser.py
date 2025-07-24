import os
import io
import csv
import docx
import google.generativeai as genai

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

SCOPES = ['https://www.googleapis.com/auth/drive']
TO_BE_PROCESSED_FOLDER_ID = 'REMOVED'

def get_google_drive_service():
  creds: None
  if os.path.exists('service_account.json'):
    creds = service_account.Credentials.from_service_account_file('service_account.json', scopes=SCOPES)
  return build('drive', 'v3', credentials=creds)

def get_file_list(service):
  full_results = []
  results = service.files().list(q=f"'{TO_BE_PROCESSED_FOLDER_ID}' in parents", pageSize=1000, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
  next_page_token = results.get('nextPageToken')
  full_results = (results.get('files', []))
  
  while (next_page_token != None):
      results = service.files().list(q=f"'{TO_BE_PROCESSED_FOLDER_ID}' in parents", pageSize=1000, pageToken=next_page_token, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
      next_page_token = results.get('nextPageToken')
      full_results.extend(results.get('files', []))
  
  return full_results

def read_word_file_from_drive(service, file_id):
  request = service.files().get_media(fileId=file_id)
  file_stream = io.BytesIO()
  downloader = MediaIoBaseDownload(file_stream, request)
  done = False
  while not done:
    done = downloader.next_chunk()
  file_stream.seek(0)
  return file_stream


def extract_text(document):
  text = []
  for section in document.sections:
    for paragraph in section.header.paragraphs:
      text.append(paragraph.text)
    for paragraph in section.footer.paragraphs:
      text.append(paragraph.text)
  for paragraph in document.paragraphs:
    text.append(paragraph.text)
  for table in document.tables:
    for row in table.rows:
      for cell in row.cells:
        text.append(cell.text)
  return text
  
    
def parse_documents(file_list, service):
  for file in file_list:
      file_id = file['id']
      file_name = file['name']
      word_stream = read_word_file_from_drive(service, file_id)
      document = docx.Document(word_stream)
      document_text = extract_text(document)
      prompt_text = "\n".join(document_text)
  return prompt_text
      

def call_gemini(prompt_text):
  genai.configure(api_key='REMOVED')
  model = genai.GenerativeModel('gemini-2.5-flash')
  response = model.generate_content(
    "Task: You will be given the full text of an agreement. Extract every milestone obligation and output them only as CSV rows with the following six columns, in this order:\n"
    "AgreementNumber — the value that follows the literal text “OTC Agreement Number:” (e.g., A-2025-2025).\n"
    "MilestoneName — one of exactly two strings:\n"
    "- PERFORMANCE MILESTONES for obligations listed in Section 9 (Performance Milestones) or described with phrasing like “Licensee will complete milestone ,,,”.\n"
    "- MILESTONE PAYMENTS for obligations listed in Section 10.5 (Milestone Payments) or described with phrasing like 'Upon ...' or 'Licensee shall pay ...'.\n"
    "MilestoneTargetDate — the due date for the obligation in YYYY-MM-DD format.\n"
    "- If the milestone has no explicit date, leave this field blank.\n"
    "MilestoneDescription — For each milestone (e.g., a, b, c), include ALL sentences in that item, from the first sentence starting with 'Upon ...' or 'Licensee shall pay ...' to the last sentence **before the next item or section**. Wrap the entire block in double quotes. Do not shorten or stop early.\n"
    "- These will only ever be found inside sections titled 'Performance Milestones' and 'Milestone Payments'.\n"
    "- If the final character in the milestone description is a semicolon, remove it. Do not include trailing semicolons.\n"
    "MilestoneSetDeadline — TRUE if a due date exists for the milestone, otherwise FALSE.\n"
    "MilestonePayment — the numeric dollar amount associated with the milestone, with no currency symbol or commas (e.g., 3610).\n"
    "- If no payment is specified, leave this field blank.\n"
    "Formatting rules:\n"
    "- Output only CSV text.\n"
    "- Use a single comma (\",\") as the delimiter between each field.\n"
    "- Preserve the order in which milestones appear in the agreement.\n"
    "- Do NOT include a header row. Begin immediately with the first milestone row.\n"
    "Input: " + prompt_text
  , generation_config={"temperature": 0})
  return response
    
def main():
  service = get_google_drive_service()
  file_list = get_file_list(service)
      
  if not file_list:
    print('No files found in "To Be Processed" folder.')
    return

  prompt_text = parse_documents(file_list, service)
  output_raw_text = call_gemini(prompt_text)
  output_text = ''.join(part.text for part in output_raw_text.candidates[0].content.parts)
  file_object = io.StringIO(output_text)
  with open('output.csv', 'w', newline='') as csvfile:
    writer = csv.writer(csvfile)
    reader = csv.reader(file_object)
    for row in reader:
      writer.writerow(row)
    
    
if __name__ == '__main__':
    main()